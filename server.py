from __future__ import annotations

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

import os
import uuid
import tempfile

import fitz  # PyMuPDF

# python-docx: Document() es una "factory", no la clase para typing
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.shared import Inches

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/", response_class=HTMLResponse)
def home():
    with open("templates/index.html", "r", encoding="utf-8") as f:
        return f.read()


def pdf_first_page_to_png_bytes(pdf_bytes: bytes, dpi: int = 150) -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if doc.page_count < 1:
        raise ValueError("El PDF no tiene páginas.")
    if doc.page_count != 1:
        raise ValueError("El PDF debe ser de UNA sola página.")

    page = doc.load_page(0)
    zoom = dpi / 72.0
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return pix.tobytes("png")


def insert_image_at_placeholder(
    docx_path: str,
    image_png: bytes,
    placeholder: str = "{{PDF_IMAGE}}",
    width_inches: float = 6.5,
) -> DocxDocumentType:
    """
    Busca placeholder en:
    - body (párrafos y tablas)
    - headers/footers (párrafos y tablas)
    y lo reemplaza insertando la imagen.
    """
    doc = DocxDocument(docx_path)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
        tmp_img.write(image_png)
        tmp_img_path = tmp_img.name

    def replace_in_paragraphs(paragraphs) -> bool:
        for p in paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, "").strip()
                run = p.add_run()
                run.add_picture(tmp_img_path, width=Inches(width_inches))
                return True
        return False

    def replace_in_tables(tables) -> bool:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if replace_in_paragraphs(cell.paragraphs):
                        return True
                    if cell.tables and replace_in_tables(cell.tables):
                        return True
        return False

    try:
        # Body
        if replace_in_paragraphs(doc.paragraphs):
            return doc
        if replace_in_tables(doc.tables):
            return doc

        # Headers/footers
        for section in doc.sections:
            header = section.header
            footer = section.footer

            if replace_in_paragraphs(header.paragraphs):
                return doc
            if replace_in_tables(header.tables):
                return doc

            if replace_in_paragraphs(footer.paragraphs):
                return doc
            if replace_in_tables(footer.tables):
                return doc

        raise ValueError(f"No encontré el marcador {placeholder} en la plantilla (ni body ni header/footer).")
    finally:
        try:
            os.remove(tmp_img_path)
        except Exception:
            pass


@app.post("/generate")
async def generate(template: UploadFile = File(...), pdf: UploadFile = File(...)):
    template_name = template.filename or ""
    pdf_name = pdf.filename or ""

    if not template_name.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="La plantilla debe ser .docx")
    if not pdf_name.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="El archivo debe ser .pdf")

    template_bytes = await template.read()
    pdf_bytes = await pdf.read()

    try:
        png_bytes = pdf_first_page_to_png_bytes(pdf_bytes, dpi=150)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    out_name = f"resultado_{uuid.uuid4().hex}.docx"
    out_path = os.path.join(tempfile.gettempdir(), out_name)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(template_bytes)
        tmp_docx_path = tmp_docx.name

    try:
        doc = insert_image_at_placeholder(tmp_docx_path, png_bytes)
        doc.save(out_path)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        try:
            os.remove(tmp_docx_path)
        except Exception:
            pass

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="resultado.docx",
    )


@app.post("/insert-pdf-image")
async def insert_pdf_image(docx: UploadFile = File(...), pdf: UploadFile = File(...)):
    docx_name = docx.filename or ""
    pdf_name = pdf.filename or ""

    if not docx_name.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="El archivo docx debe ser .docx")
    if not pdf_name.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="El archivo pdf debe ser .pdf")

    docx_bytes = await docx.read()
    pdf_bytes = await pdf.read()

    try:
        png_bytes = pdf_first_page_to_png_bytes(pdf_bytes, dpi=150)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    out_name = f"resultado_{uuid.uuid4().hex}.docx"
    out_path = os.path.join(tempfile.gettempdir(), out_name)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(docx_bytes)
        tmp_docx_path = tmp_docx.name

    try:
        doc_obj = insert_image_at_placeholder(tmp_docx_path, png_bytes, placeholder="{{PDF_IMAGE}}")
        doc_obj.save(out_path)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        try:
            os.remove(tmp_docx_path)
        except Exception:
            pass

    return FileResponse(
        out_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="resultado.docx",
    )

@app.get("/healthz")
def healthz():
    return {"ok": True}
